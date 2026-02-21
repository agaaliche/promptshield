"""Document ingestion — load files of various formats and convert to page bitmaps + text."""

from __future__ import annotations

import logging
import mimetypes
import ctypes
import shutil
import subprocess
import uuid
from pathlib import Path
from typing import Optional

import pypdfium2 as pdfium
from PIL import Image

from core.config import config
from core.ocr.engine import ocr_page_image
from models.schemas import BBox, DocumentInfo, DocumentStatus, PageData, TextBlock

logger = logging.getLogger(__name__)


def _cluster_into_lines(text_blocks: list[TextBlock]) -> list[list[TextBlock]]:
    """Group text blocks into visual lines by y-proximity, then sort
    left-to-right within each line.

    Plain ``sorted(blocks, key=(y0, x0))`` can mis-order words on the
    same visual line when their y0 values differ slightly (common with
    ascenders/descenders in PDF extraction and OCR noise).  Clustering
    by y-center avoids this.
    """
    if not text_blocks:
        return []

    by_y = sorted(text_blocks, key=lambda b: (b.bbox.y0 + b.bbox.y1) / 2)

    lines: list[list[TextBlock]] = []
    cur_line: list[TextBlock] = []
    line_yc: float = 0.0
    line_h: float = 0.0

    for block in by_y:
        bh = block.bbox.y1 - block.bbox.y0
        byc = (block.bbox.y0 + block.bbox.y1) / 2

        if not cur_line:
            cur_line.append(block)
            line_yc = byc
            line_h = bh
        else:
            tolerance = max(line_h, bh) * 0.5
            if abs(byc - line_yc) <= tolerance:
                cur_line.append(block)
                line_h = max(line_h, bh)
                n = len(cur_line)
                line_yc = (line_yc * (n - 1) + byc) / n
            else:
                lines.append(sorted(cur_line, key=lambda b: b.bbox.x0))
                cur_line = [block]
                line_yc = byc
                line_h = bh

    if cur_line:
        lines.append(sorted(cur_line, key=lambda b: b.bbox.x0))

    return lines


def _build_full_text(text_blocks: list[TextBlock]) -> str:
    """Build full_text from text blocks, inserting newlines between
    lines that are spatially separated (different vertical position).

    Blocks are clustered into visual lines (tolerant of minor y-position
    differences from ascenders, descenders, or OCR noise) and sorted
    left-to-right within each line, producing correct reading order.

    Within each visual line, a large horizontal gap between consecutive
    blocks (> 3× average block height) is treated as a **column
    boundary** and a newline is inserted instead of a space.  This
    prevents regex / NER from seeing text in separate columns as one
    continuous phrase (works even for single-line columns).

    This gives NER/regex much better input than a flat space-joined
    string because sentence boundaries are preserved.
    """
    if not text_blocks:
        return ""

    lines = _cluster_into_lines(text_blocks)

    parts: list[str] = []
    prev_y: float | None = None
    line_height = 0.0

    for line_blocks in lines:
        line_top = min(b.bbox.y0 for b in line_blocks)
        lh = max(b.bbox.y1 for b in line_blocks) - line_top

        # Column-gap threshold: a horizontal gap > 3× average block
        # height on this line indicates a column boundary.
        avg_h = (
            sum(b.bbox.y1 - b.bbox.y0 for b in line_blocks)
            / len(line_blocks)
        )
        col_gap_threshold = max(avg_h * 3, 15.0)

        for i, block in enumerate(line_blocks):
            if prev_y is not None or i > 0:
                if i == 0:
                    gap = line_top - prev_y if prev_y is not None else 0
                    if line_height > 0 and gap > line_height * 0.6:
                        parts.append("\n")
                    else:
                        parts.append(" ")
                else:
                    # Detect column gap within the same visual line
                    prev_block = line_blocks[i - 1]
                    h_gap = block.bbox.x0 - prev_block.bbox.x1
                    if h_gap > col_gap_threshold:
                        parts.append("\n")
                    else:
                        parts.append(" ")
            parts.append(block.text)

        prev_y = line_top
        line_height = lh

    return "".join(parts).strip()


# Supported MIME types
PDF_TYPES = {"application/pdf"}
IMAGE_TYPES = {"image/jpeg", "image/png", "image/tiff", "image/bmp", "image/webp"}
OFFICE_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",   # docx
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",          # xlsx
    "application/vnd.openxmlformats-officedocument.presentationml.presentation",  # pptx
    "application/msword",
    "application/vnd.ms-excel",
    "application/vnd.ms-powerpoint",
}

SUPPORTED_EXTENSIONS = {
    ".pdf", ".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp",
    ".docx", ".xlsx", ".pptx", ".doc", ".xls", ".ppt",
}


def guess_mime(filepath: Path) -> str:
    """Guess MIME type from file extension."""
    mime, _ = mimetypes.guess_type(str(filepath))
    return mime or "application/octet-stream"


def _libreoffice_convert_to_pdf(src: Path, out_dir: Path) -> Path:
    """Convert an Office document to PDF using LibreOffice headless."""
    lo_cmd = shutil.which("libreoffice") or shutil.which("soffice")
    if lo_cmd is None:
        raise RuntimeError(
            "LibreOffice is required to process Office documents but was not found on PATH. "
            "Install LibreOffice and ensure 'libreoffice' or 'soffice' is available."
        )
    subprocess.run(
        [lo_cmd, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(src)],
        check=True,
        capture_output=True,
        timeout=120,
    )
    pdf_path = out_dir / f"{src.stem}.pdf"
    if not pdf_path.exists():
        raise RuntimeError(f"LibreOffice conversion failed — expected output at {pdf_path}")
    return pdf_path


def _docx_to_pdf(src: Path, out_dir: Path) -> Path:
    """Convert a DOCX file to PDF using python-docx and reportlab.

    This is a pure-Python alternative to LibreOffice for docx files.
    Extracts paragraphs, tables, headers/footers, and renders them into
    a PDF preserving text content and basic formatting for accurate
    text extraction and PII detection.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from xml.sax.saxutils import escape

    out_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = out_dir / f"{src.stem}.pdf"

    docx_doc = Document(str(src))

    # ── styles ──────────────────────────────────────────────────────
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "DocxBody", parent=styles["Normal"],
        fontSize=10, leading=13, spaceAfter=4,
    )
    bold_style = ParagraphStyle(
        "DocxBold", parent=body_style,
        fontName="Helvetica-Bold",
    )
    heading_styles: dict[int, ParagraphStyle] = {}
    for level, (fs, lead) in enumerate(
        [(16, 20), (14, 18), (12, 16), (11, 15), (10, 14)], start=1
    ):
        heading_styles[level] = ParagraphStyle(
            f"DocxH{level}", parent=styles["Normal"],
            fontSize=fs, leading=lead, spaceAfter=6, spaceBefore=8,
            fontName="Helvetica-Bold",
        )
    cell_style = ParagraphStyle(
        "DocxCell", parent=styles["Normal"],
        fontSize=9, leading=11, wordWrap="CJK",
    )

    # ── helper: convert a docx paragraph to reportlab Paragraph ─────
    def _para_to_flowable(para):
        """Return a reportlab Paragraph (or None for empty paragraphs)."""
        text_parts: list[str] = []
        for run in para.runs:
            t = escape(run.text) if run.text else ""
            if not t:
                continue
            # Preserve basic formatting via HTML-like tags
            if run.bold and run.italic:
                t = f"<b><i>{t}</i></b>"
            elif run.bold:
                t = f"<b>{t}</b>"
            elif run.italic:
                t = f"<i>{t}</i>"
            if run.underline:
                t = f"<u>{t}</u>"
            text_parts.append(t)

        full = "".join(text_parts).strip()
        if not full:
            return None

        # Pick style based on heading level
        style_name = para.style.name if para.style else ""
        lvl = 0
        if style_name.startswith("Heading"):
            try:
                lvl = int(style_name.split()[-1])
            except (ValueError, IndexError):
                lvl = 1
        pstyle = heading_styles.get(lvl, body_style)

        # Alignment
        align = para.alignment
        if align == WD_ALIGN_PARAGRAPH.CENTER:
            pstyle = ParagraphStyle(f"_c{id(para)}", parent=pstyle, alignment=1)
        elif align == WD_ALIGN_PARAGRAPH.RIGHT:
            pstyle = ParagraphStyle(f"_r{id(para)}", parent=pstyle, alignment=2)

        return Paragraph(full, pstyle)

    # ── helper: convert a docx table to reportlab Table ─────────────
    def _table_to_flowable(tbl):
        """Return a reportlab Table flowable."""
        data: list[list] = []
        for row in tbl.rows:
            row_data: list = []
            for cell in row.cells:
                cell_text = escape(cell.text.strip()) if cell.text else ""
                if len(cell_text) > 200:
                    cell_text = cell_text[:197] + "..."
                row_data.append(Paragraph(cell_text, cell_style))
            data.append(row_data)
        if not data:
            return None
        ncols = max(len(r) for r in data)
        available = letter[0] - 1 * inch
        col_w = min(available / max(ncols, 1), 2.5 * inch)
        t = Table(data, colWidths=[col_w] * ncols)
        t.setStyle(TableStyle([
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        return t

    # ── walk the document body in order ─────────────────────────────
    # python-docx exposes paragraphs and tables as separate lists, but
    # the underlying XML has them interleaved.  Iterate the body XML
    # children to preserve correct ordering.
    from docx.oxml.ns import qn

    elements: list = []

    # Header text (first section only, for PII detection)
    for section in docx_doc.sections:
        for hdr in (section.header,):
            if hdr and not hdr.is_linked_to_previous:
                for p in hdr.paragraphs:
                    fl = _para_to_flowable(p)
                    if fl:
                        elements.append(fl)
        break  # first section headers only

    body = docx_doc.element.body
    for child in body:
        tag = child.tag
        if tag == qn("w:p"):
            # paragraph
            para = None
            for p in docx_doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para is None:
                continue
            # Page break detection
            if para.style and para.style.name and "page" in para.style.name.lower():
                elements.append(PageBreak())
            fl = _para_to_flowable(para)
            if fl:
                elements.append(fl)
            else:
                elements.append(Spacer(1, 4))
        elif tag == qn("w:tbl"):
            tbl = None
            for t in docx_doc.tables:
                if t._element is child:
                    tbl = t
                    break
            if tbl:
                fl = _table_to_flowable(tbl)
                if fl:
                    elements.append(Spacer(1, 4))
                    elements.append(fl)
                    elements.append(Spacer(1, 4))
        elif tag == qn("w:sectPr"):
            # Section break — treat as page break
            elements.append(PageBreak())

    # Footer text (first section only)
    for section in docx_doc.sections:
        for ftr in (section.footer,):
            if ftr and not ftr.is_linked_to_previous:
                for p in ftr.paragraphs:
                    fl = _para_to_flowable(p)
                    if fl:
                        elements.append(fl)
        break

    # ── build PDF ───────────────────────────────────────────────────
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    if not elements:
        elements.append(Paragraph("(Empty document)", styles["Normal"]))

    doc.build(elements)

    if not pdf_path.exists():
        raise RuntimeError(
            f"DOCX to PDF conversion failed — expected output at {pdf_path}"
        )

    logger.info(f"Converted DOCX to PDF: {pdf_path}")
    return pdf_path


def _xlsx_to_pdf(src: Path, out_dir: Path) -> Path:
    """Convert an Excel xlsx file to PDF using openpyxl and reportlab.
    
    This is a pure-Python alternative to LibreOffice for xlsx files.
    Creates a PDF with each sheet as a separate page, preserving cell
    content in a tabular layout for accurate text extraction.
    """
    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    
    out_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = out_dir / f"{src.stem}.pdf"
    
    # Load workbook
    wb = load_workbook(src, data_only=True)  # data_only=True to get calculated values
    
    # Create PDF document with landscape orientation for spreadsheets
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=landscape(letter),
        leftMargin=0.5*inch,
        rightMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch,
    )
    
    styles = getSampleStyleSheet()
    # Create a style for cell content that handles wrapping
    cell_style = ParagraphStyle(
        'CellStyle',
        parent=styles['Normal'],
        fontSize=8,
        leading=10,
        wordWrap='CJK',  # Better wrapping for all text
    )
    header_style = ParagraphStyle(
        'HeaderStyle',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=12,
    )
    
    elements = []
    
    for sheet_idx, sheet_name in enumerate(wb.sheetnames):
        ws = wb[sheet_name]
        
        # Add sheet name as header
        elements.append(Paragraph(f"Sheet: {sheet_name}", header_style))
        elements.append(Spacer(1, 0.1*inch))
        
        # Find the used range
        if ws.max_row is None or ws.max_column is None:
            continue
            
        max_row = min(ws.max_row, 500)  # Limit rows to prevent huge PDFs
        max_col = min(ws.max_column, 26)  # Limit columns (A-Z)
        
        if max_row == 0 or max_col == 0:
            elements.append(Paragraph("(Empty sheet)", styles['Normal']))
            if sheet_idx < len(wb.sheetnames) - 1:
                elements.append(PageBreak())
            continue
        
        # Build table data
        table_data = []
        for row_idx in range(1, max_row + 1):
            row_data = []
            for col_idx in range(1, max_col + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                value = cell.value
                if value is None:
                    value = ""
                elif isinstance(value, (int, float)):
                    # Format numbers nicely
                    if isinstance(value, float) and value == int(value):
                        value = str(int(value))
                    else:
                        value = str(value)
                else:
                    value = str(value)
                # Truncate very long cell values
                if len(value) > 100:
                    value = value[:97] + "..."
                # Wrap text in Paragraph for proper handling
                row_data.append(Paragraph(value, cell_style))
            table_data.append(row_data)
        
        if not table_data:
            elements.append(Paragraph("(Empty sheet)", styles['Normal']))
            if sheet_idx < len(wb.sheetnames) - 1:
                elements.append(PageBreak())
            continue
        
        # Calculate column widths (distribute evenly with max width)
        available_width = landscape(letter)[0] - 1*inch  # Page width minus margins
        col_width = min(available_width / max_col, 1.5*inch)
        col_widths = [col_width] * max_col
        
        # Create table
        table = Table(table_data, colWidths=col_widths, repeatRows=1)
        table.setStyle(TableStyle([
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),  # Header row
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 3),
            ('RIGHTPADDING', (0, 0), (-1, -1), 3),
            ('TOPPADDING', (0, 0), (-1, -1), 2),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ]))
        
        elements.append(table)
        
        # Add page break between sheets
        if sheet_idx < len(wb.sheetnames) - 1:
            elements.append(PageBreak())
    
    # Build PDF
    if elements:
        doc.build(elements)
    else:
        # Create empty PDF with message
        elements.append(Paragraph("(Empty workbook)", styles['Normal']))
        doc.build(elements)
    
    wb.close()
    
    if not pdf_path.exists():
        raise RuntimeError(f"xlsx to PDF conversion failed — expected output at {pdf_path}")
    
    logger.info(f"Converted xlsx to PDF: {pdf_path}")
    return pdf_path


def _is_rotated_word(char_y_centers: list[float], char_heights: list[float]) -> bool:
    """Return True if the accumulated character positions indicate rotated text.

    For horizontal text, all characters share roughly the same y-centre.
    For rotated/diagonal text (watermarks), the y-centres of successive
    characters shift significantly — the vertical spread of centres will
    exceed a fraction of the average character height.

    Punctuation marks (periods, commas, apostrophes) and accent marks have
    much smaller bounding boxes positioned at different baselines than the
    main letter glyphs.  Without filtering them out, a word like ``B.N.``
    or ``l'exercice`` would look "rotated" even though it's perfectly
    horizontal.  We therefore exclude characters whose height is less than
    70 % of the median character height before computing the spread.
    """
    if len(char_y_centers) < 2:
        return False

    # --- filter out small characters (punctuation / accents / superscripts) ---
    sorted_h = sorted(char_heights)
    n = len(sorted_h)
    median_h = sorted_h[n // 2] if n % 2 == 1 else (sorted_h[n // 2 - 1] + sorted_h[n // 2]) / 2.0
    height_threshold = median_h * 0.70

    filtered_yc: list[float] = []
    filtered_h: list[float] = []
    for yc, h in zip(char_y_centers, char_heights):
        if h >= height_threshold:
            filtered_yc.append(yc)
            filtered_h.append(h)

    if len(filtered_yc) < 2:
        return False

    y_spread = max(filtered_yc) - min(filtered_yc)
    avg_h = sum(filtered_h) / len(filtered_h)
    # Horizontal text: y_spread ≈ 0.  45° rotated: y_spread ≈ word width.
    # Threshold raised to 0.65 to tolerate accented capitals (É, Ô, …) whose
    # taller bounding-box shifts the y-centre slightly.
    return y_spread > avg_h * 0.65


def _extract_text_blocks_from_page(pdf_page: pdfium.PdfPage, page_index: int) -> list[TextBlock]:
    """Extract word-level text blocks with bounding boxes from a PDF page.

    Rotated text (diagonal watermarks, etc.) is automatically discarded.
    Font weight and italic flags are extracted from pdfium per-character
    metadata and propagated to the resulting TextBlock.
    """
    textpage = pdf_page.get_textpage()
    full_text = textpage.get_text_range()

    blocks: list[TextBlock] = []
    if not full_text.strip():
        return blocks

    # Use character-level extraction and group into words
    n_chars = textpage.count_chars()
    if n_chars == 0:
        return blocks

    # ── Helper: query font weight / italic / name / size for a character index ───
    _raw_tp = textpage.raw  # underlying FPDF_TEXTPAGE handle
    _fw_func = pdfium.raw.FPDFText_GetFontWeight
    _fi_func = pdfium.raw.FPDFText_GetFontInfo
    _fs_func = pdfium.raw.FPDFText_GetFontSize
    _fi_buf = ctypes.create_string_buffer(256)
    _fi_flags = ctypes.c_int(0)

    def _char_font_style(idx: int) -> tuple[bool, bool, str, float]:
        """Return (is_bold, is_italic, font_name, font_size) for character *idx*."""
        weight = _fw_func(_raw_tp, idx)
        bold = weight >= 700
        _fi_flags.value = 0
        _fi_func(
            _raw_tp, idx,
            ctypes.cast(_fi_buf, ctypes.c_void_p),
            ctypes.c_ulong(256),
            ctypes.byref(_fi_flags),
        )
        italic = bool(_fi_flags.value & 0x01)
        font_name = _fi_buf.value.decode("utf-8", errors="replace").strip()
        font_size = float(_fs_func(_raw_tp, idx))
        return bold, italic, font_name, font_size

    current_word = ""
    word_x0 = word_y0 = word_x1 = word_y1 = 0.0
    word_start_idx = 0
    word_index = 0
    # Per-character tracking for rotation detection
    char_y_centers: list[float] = []
    char_heights: list[float] = []
    rotated_skipped = 0
    # Font-style accumulators for the current word
    word_bold_votes = 0
    word_italic_votes = 0
    word_char_count = 0
    # Font name / size — sampled from the first character of each word
    word_font_name: str = ""
    word_font_size: float = 0.0

    for i in range(n_chars):
        char = textpage.get_text_range(index=i, count=1)
        charbox = textpage.get_charbox(i)  # (left, bottom, right, top) in PDF coords

        if char.strip() == "":
            # End of word — flush if we have accumulated text
            if current_word:
                if _is_rotated_word(char_y_centers, char_heights):
                    rotated_skipped += 1
                else:
                    # pypdfium2 charbox is (left, bottom, right, top) in PDF coords
                    # Convert to top-left origin: y0 = page_height - top, y1 = page_height - bottom
                    page_height = pdf_page.get_height()
                    # Majority-vote for font style across characters in the word
                    w_bold = word_bold_votes > word_char_count / 2
                    w_italic = word_italic_votes > word_char_count / 2
                    blocks.append(TextBlock(
                        text=current_word,
                        bbox=BBox(
                            x0=word_x0,
                            y0=page_height - word_y1,  # top in screen coords
                            x1=word_x1,
                            y1=page_height - word_y0,  # bottom in screen coords
                        ),
                        confidence=1.0,
                        block_index=0,
                        line_index=0,
                        word_index=word_index,
                        is_ocr=False,
                        is_bold=w_bold,
                        is_italic=w_italic,
                        font_size=word_font_size,
                        font_family=word_font_name,
                    ))
                word_index += 1
                current_word = ""
                char_y_centers = []
                char_heights = []
                word_bold_votes = 0
                word_italic_votes = 0
                word_char_count = 0
                word_font_name = ""
                word_font_size = 0.0
        else:
            left, bottom, right, top = charbox
            if not current_word:
                word_x0 = left
                word_y0 = bottom
                word_x1 = right
                word_y1 = top
                word_start_idx = i
            else:
                word_x0 = min(word_x0, left)
                word_y0 = min(word_y0, bottom)
                word_x1 = max(word_x1, right)
                word_y1 = max(word_y1, top)
            current_word += char
            char_y_centers.append((bottom + top) / 2.0)
            char_heights.append(top - bottom)
            # Track font style (sample first char for name/size)
            try:
                cb, ci, fname, fsize = _char_font_style(i)
                word_bold_votes += int(cb)
                word_italic_votes += int(ci)
                if word_char_count == 0:  # first char of word
                    word_font_name = fname
                    word_font_size = fsize
            except Exception:
                pass  # pdfium may fail on some exotic fonts
            word_char_count += 1

    # Flush last word
    if current_word:
        if _is_rotated_word(char_y_centers, char_heights):
            rotated_skipped += 1
        else:
            page_height = pdf_page.get_height()
            w_bold = word_bold_votes > word_char_count / 2 if word_char_count else False
            w_italic = word_italic_votes > word_char_count / 2 if word_char_count else False
            blocks.append(TextBlock(
                text=current_word,
                bbox=BBox(
                    x0=word_x0,
                    y0=page_height - word_y1,
                    x1=word_x1,
                    y1=page_height - word_y0,
                ),
                confidence=1.0,
                block_index=0,
                line_index=0,
                word_index=word_index,
                is_ocr=False,
                is_bold=w_bold,
                is_italic=w_italic,
                font_size=word_font_size,
                font_family=word_font_name,
            ))

    if rotated_skipped:
        import logging
        logging.getLogger(__name__).info(
            f"Page {page_index + 1}: discarded {rotated_skipped} rotated "
            f"text block(s) (watermarks/diagonal text)"
        )

    return blocks


def _render_page_bitmap(pdf_page: pdfium.PdfPage, page_index: int, doc_id: str) -> Path:
    """Render a PDF page to a PNG bitmap."""
    scale = config.render_dpi / 72  # PDF default is 72 DPI
    bitmap = pdf_page.render(scale=scale)
    pil_image = bitmap.to_pil()

    out_path = config.temp_dir / doc_id / f"page_{page_index + 1:04d}.png"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    pil_image.save(str(out_path), "PNG")
    pil_image.close()
    del bitmap
    return out_path


def _has_embedded_images(pdf_page: pdfium.PdfPage, min_area_fraction: float = 0.005) -> bool:
    """Check if page has embedded images that may contain text.
    
    Returns True if the page contains image objects covering at least
    min_area_fraction of the page area. This indicates potential text-in-image
    content that requires OCR for proper extraction.
    
    Args:
        pdf_page: The PDF page to check
        min_area_fraction: Minimum fraction of page area for an image to count (default 0.5%)
    """
    page_width = pdf_page.get_width()
    page_height = pdf_page.get_height()
    page_area = page_width * page_height
    min_area = page_area * min_area_fraction
    
    try:
        for obj in pdf_page.get_objects():
            if type(obj).__name__ == 'PdfImage':
                bounds = obj.get_bounds()
                left, bottom, right, top = bounds
                img_area = (right - left) * (top - bottom)
                if img_area >= min_area:
                    return True
    except Exception as e:
        logger.debug(f"Error checking for images: {e}")
    
    return False


def _merge_ocr_blocks(existing: list[TextBlock], ocr_blocks: list[TextBlock]) -> list[TextBlock]:
    """Merge OCR blocks with existing text blocks, avoiding duplicates.
    
    Only adds OCR blocks that don't significantly overlap with existing text.
    """
    if not ocr_blocks:
        return existing
    if not existing:
        return ocr_blocks
    
    def overlaps(b1: BBox, b2: BBox) -> bool:
        """Check if two bboxes overlap significantly."""
        # Compute intersection
        ix0 = max(b1.x0, b2.x0)
        iy0 = max(b1.y0, b2.y0)
        ix1 = min(b1.x1, b2.x1)
        iy1 = min(b1.y1, b2.y1)
        
        if ix1 <= ix0 or iy1 <= iy0:
            return False
        
        # Check if intersection is >50% of the smaller box
        inter_area = (ix1 - ix0) * (iy1 - iy0)
        b1_area = max((b1.x1 - b1.x0) * (b1.y1 - b1.y0), 1)
        b2_area = max((b2.x1 - b2.x0) * (b2.y1 - b2.y0), 1)
        min_area = min(b1_area, b2_area)
        
        return inter_area > 0.5 * min_area
    
    merged = list(existing)
    for ocr_b in ocr_blocks:
        # Only add if no significant overlap with existing blocks
        if not any(overlaps(ocr_b.bbox, e.bbox) for e in existing):
            merged.append(ocr_b)
    
    return merged


# Type alias for progress callbacks
ProgressCallback = Optional[callable]


def _process_pdf(
    pdf_path: Path, 
    doc_id: str, 
    progress_callback: ProgressCallback = None
) -> list[PageData]:
    """Process a PDF file into page data.

    **Phase 1 — sequential (PDFium):** render bitmaps and extract text
    for every page using a single ``PdfDocument`` handle.  PDFium's C
    library is *not* thread-safe (concurrent handles cause heap
    corruption / native breakpoint crashes on Windows).

    **Phase 2 — parallel OCR (Tesseract):** pages whose embedded text is
    too sparse are sent to OCR in parallel threads.  Tesseract is a
    separate process per invocation and fully thread-safe, so this
    recovers most of the I1 parallelism for scanned/image-based PDFs.
    
    Args:
        pdf_path: Path to the PDF file.
        doc_id: Unique document identifier.
        progress_callback: Optional callback for progress updates.
            Called with (phase, current_page, total_pages, ocr_done, ocr_total, message).
    """
    def _report(phase: str, current: int, total: int, ocr_done: int = 0, ocr_total: int = 0, message: str = ""):
        if progress_callback:
            try:
                progress_callback(phase, current, total, ocr_done, ocr_total, message)
            except Exception:
                pass  # Don't let callback errors break processing
    
    # ── Phase 1: sequential PDFium extraction ──────────────────────
    doc = pdfium.PdfDocument(str(pdf_path))
    try:
        n_pages = len(doc)
        if n_pages == 0:
            return []
        
        _report("extracting", 0, n_pages, 0, 0, f"Loading {n_pages} pages...")

        pages: list[PageData] = []
        ocr_needed: list[int] = []  # indices into *pages* that need OCR

        for page_index in range(n_pages):
            pdf_page = doc[page_index]
            width = pdf_page.get_width()
            height = pdf_page.get_height()

            bitmap_path = _render_page_bitmap(pdf_page, page_index, doc_id)
            text_blocks = _extract_text_blocks_from_page(pdf_page, page_index)
            full_text = _build_full_text(text_blocks)

            pages.append(PageData(
                page_number=page_index + 1,
                width=width,
                height=height,
                bitmap_path=str(bitmap_path),
                text_blocks=text_blocks,
                full_text=full_text,
            ))
            
            # Report extraction progress
            _report("extracting", page_index + 1, n_pages, 0, 0, f"Extracting page {page_index + 1} of {n_pages}...")

            # OCR needed if: (a) sparse text, or (b) page has embedded images
            if len(full_text.strip()) < 20:
                ocr_needed.append(page_index)
            elif _has_embedded_images(pdf_page):
                logger.info(f"Page {page_index + 1}: embedded images detected, will run hybrid OCR")
                ocr_needed.append(page_index)

    finally:
        doc.close()

    # ── Phase 2: parallel OCR for sparse-text pages ────────────────
    if not ocr_needed:
        _report("complete", n_pages, n_pages, 0, 0, "Processing complete")
        return pages
    
    ocr_total = len(ocr_needed)
    ocr_done = [0]  # Use list to allow mutation in nested function
    
    _report("ocr", n_pages, n_pages, 0, ocr_total, f"Running OCR on {ocr_total} pages...")

    def _ocr_one(idx: int) -> tuple[int, list[TextBlock], bool]:
        """Returns (index, ocr_blocks, should_merge).
        
        should_merge=True when page already has content (hybrid mode).
        """
        p = pages[idx]
        has_existing_content = len(p.text_blocks) >= 10
        if has_existing_content:
            logger.info(f"Page {p.page_number}: hybrid OCR (has {len(p.text_blocks)} text blocks + images)")
        else:
            logger.info(f"Page {p.page_number}: sparse text ({len(p.full_text)} chars), running OCR")
        blocks = ocr_page_image(Path(p.bitmap_path), p.width, p.height)
        return idx, blocks, has_existing_content

    def _apply_ocr_result(idx: int, ocr_blocks: list[TextBlock], should_merge: bool):
        if not ocr_blocks:
            return
        if should_merge:
            # Merge OCR blocks with existing text (hybrid mode)
            merged = _merge_ocr_blocks(pages[idx].text_blocks, ocr_blocks)
            pages[idx] = pages[idx].model_copy(update={
                "text_blocks": merged,
                "full_text": _build_full_text(merged),
            })
            logger.info(f"Page {pages[idx].page_number}: merged {len(ocr_blocks)} OCR blocks, now {len(merged)} total")
        else:
            # Replace entirely (sparse text case)
            pages[idx] = pages[idx].model_copy(update={
                "text_blocks": ocr_blocks,
                "full_text": _build_full_text(ocr_blocks),
            })

    if len(ocr_needed) == 1:
        # Single page — skip thread overhead
        idx, ocr_blocks, should_merge = _ocr_one(ocr_needed[0])
        _apply_ocr_result(idx, ocr_blocks, should_merge)
        ocr_done[0] = 1
        _report("ocr", n_pages, n_pages, 1, ocr_total, f"OCR complete (1/{ocr_total} pages)")
    else:
        import os
        from concurrent.futures import ThreadPoolExecutor, as_completed

        workers = min(4, os.cpu_count() or 2, len(ocr_needed))
        with ThreadPoolExecutor(max_workers=workers) as pool:
            futures = {pool.submit(_ocr_one, i): i for i in ocr_needed}
            for fut in as_completed(futures):
                idx, ocr_blocks, should_merge = fut.result()
                _apply_ocr_result(idx, ocr_blocks, should_merge)
                ocr_done[0] += 1
                _report("ocr", n_pages, n_pages, ocr_done[0], ocr_total, 
                       f"OCR progress: {ocr_done[0]}/{ocr_total} pages")
    
    _report("complete", n_pages, n_pages, ocr_total, ocr_total, "Processing complete")
    return pages


def _process_image(
    image_path: Path, 
    doc_id: str,
    progress_callback: ProgressCallback = None
) -> list[PageData]:
    """Process a standalone image file (always OCR)."""
    def _report(phase: str, current: int, total: int, ocr_done: int = 0, ocr_total: int = 0, message: str = ""):
        if progress_callback:
            try:
                progress_callback(phase, current, total, ocr_done, ocr_total, message)
            except Exception:
                pass
    
    _report("extracting", 0, 1, 0, 1, "Loading image...")
    
    img = Image.open(image_path)
    try:
        width, height = img.size

        # Save as PNG in temp
        out_path = config.temp_dir / doc_id / "page_0001.png"
        out_path.parent.mkdir(parents=True, exist_ok=True)

        # Upscale small images to improve OCR accuracy
        min_dim = min(width, height)
        if min_dim < 1500:
            scale = max(2, 1500 // min_dim)
            big = img.resize((width * scale, height * scale), Image.LANCZOS)
            big.save(str(out_path), "PNG")
            big.close()
            logger.info(f"Upscaled image {width}x{height} by {scale}x for OCR")
        else:
            img.save(str(out_path), "PNG")

        _report("ocr", 1, 1, 0, 1, "Running OCR on image...")
        
        # OCR is required for images
        text_blocks = ocr_page_image(out_path, float(width), float(height))
        full_text = _build_full_text(text_blocks)
        
        _report("complete", 1, 1, 1, 1, "Processing complete")

        return [PageData(
            page_number=1,
            width=float(width),
            height=float(height),
            bitmap_path=str(out_path),
            text_blocks=text_blocks,
            full_text=full_text,
        )]
    finally:
        img.close()


async def ingest_document(
    file_path: Path,
    original_filename: str,
    mime_type: Optional[str] = None,
    progress_callback: ProgressCallback = None,
) -> DocumentInfo:
    """
    Main entry point: ingest a document file, convert to bitmaps, extract text.

    Returns a DocumentInfo with pages populated.
    
    Args:
        file_path: Path to the file to ingest.
        original_filename: Original name of the uploaded file.
        mime_type: Optional MIME type (auto-detected if not provided).
        progress_callback: Optional callback for progress updates.
            Called with (phase, current_page, total_pages, ocr_done, ocr_total, message).
    """
    import asyncio

    doc_id = uuid.uuid4().hex[:12]
    if mime_type is None:
        mime_type = guess_mime(file_path)

    logger.info(f"Ingesting '{original_filename}' (mime={mime_type}, id={doc_id})")

    doc = DocumentInfo(
        doc_id=doc_id,
        original_filename=original_filename,
        file_path=str(file_path),
        mime_type=mime_type,
        status=DocumentStatus.PROCESSING,
    )

    def _do_ingest() -> list:
        if mime_type in PDF_TYPES:
            return _process_pdf(file_path, doc_id, progress_callback)
        elif mime_type in IMAGE_TYPES:
            return _process_image(file_path, doc_id, progress_callback)
        elif mime_type in OFFICE_TYPES:
            # Use native converters where possible (no LibreOffice needed)
            xlsx_mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if mime_type == xlsx_mime:
                pdf_path = _xlsx_to_pdf(file_path, config.temp_dir / doc_id)
            elif mime_type == docx_mime:
                pdf_path = _docx_to_pdf(file_path, config.temp_dir / doc_id)
            else:
                pdf_path = _libreoffice_convert_to_pdf(
                    file_path,
                    config.temp_dir / doc_id,
                )
            return _process_pdf(pdf_path, doc_id, progress_callback)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")

    try:
        doc.pages = await asyncio.to_thread(_do_ingest)

        doc.page_count = len(doc.pages)
        doc.status = DocumentStatus.EXTRACTED
        logger.info(f"Ingested {doc.page_count} pages from '{original_filename}'")

    except Exception as e:
        logger.exception(f"Failed to ingest '{original_filename}'")
        doc.status = DocumentStatus.ERROR
        raise

    return doc
