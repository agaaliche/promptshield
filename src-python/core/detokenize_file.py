"""File-based de-tokenization — replace tokens in .pdf, .docx, .xlsx, .txt files.

Reads the file, runs vault.resolve_all_tokens() on extracted text,
writes a new file with tokens replaced, and returns its path.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path

from typing import Callable

import fitz  # PyMuPDF

from core.config import config

logger = logging.getLogger(__name__)

# Supported MIME → extension mapping
_MIME_MAP: dict[str, str] = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/msword": ".doc",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
    "application/vnd.ms-excel": ".xls",
    "text/plain": ".txt",
    "text/csv": ".csv",
}


def _suffix_from_filename(filename: str) -> str:
    return Path(filename).suffix.lower()


# ---------------------------------------------------------------------------
# Per-format handlers
# ---------------------------------------------------------------------------

def _detokenize_txt(data: bytes, vault) -> tuple[bytes, int, list[str]]:
    """Plain text / CSV — straightforward string replace."""
    text = data.decode("utf-8", errors="replace")
    result, count, unresolved = vault.resolve_all_tokens(text)
    return result.encode("utf-8"), count, unresolved


def _detokenize_docx(data: bytes, vault) -> tuple[bytes, int, list[str]]:
    """DOCX — iterate over paragraphs, tables, headers/footers."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    total_replaced = 0
    all_unresolved: list[str] = []

    def _process(text: str) -> str:
        nonlocal total_replaced, all_unresolved
        result, count, unresolved = vault.resolve_all_tokens(text)
        total_replaced += count
        all_unresolved.extend(unresolved)
        return result

    # Paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = _process(run.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = _process(run.text)

    # Headers & footers
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer and header_footer.is_linked_to_previous is False:
                for para in header_footer.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = _process(run.text)

    buf = io.BytesIO()
    doc.save(buf)
    # Deduplicate unresolved
    return buf.getvalue(), total_replaced, list(set(all_unresolved))


def _detokenize_xlsx(data: bytes, vault) -> tuple[bytes, int, list[str]]:
    """XLSX — iterate over all cells in all sheets."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data))
    total_replaced = 0
    all_unresolved: list[str] = []

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value:
                    result, count, unresolved = vault.resolve_all_tokens(cell.value)
                    if count > 0:
                        cell.value = result
                        total_replaced += count
                        all_unresolved.extend(unresolved)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue(), total_replaced, list(set(all_unresolved))


def _detokenize_pdf(data: bytes, vault) -> tuple[bytes, int, list[str]]:
    """PDF — in-place token replacement preserving original layout.

    Strategy per page:
      1. Search for each token string → get pixel rects.
      2. Extract the token's visual style (font, size, color, baseline).
      3. Add erase-only redactions (white fill, no text).
      4. apply_redactions() wipes the token text.
      5. insert_text() at the original baseline with the original value
         using the extracted style — unconstrained by rect, so the
         (usually longer) original text renders at the correct size.
    """
    import re

    pdf_doc = fitz.open(stream=data, filetype="pdf")

    total_replaced = 0
    all_unresolved: list[str] = []

    try:
        # First pass: collect all unique token strings from the document
        full_text = ""
        for page in pdf_doc:
            full_text += page.get_text() + "\n"

        # Find token-shaped strings — supports both compact [P38291] and
        # legacy [ANON_TYPE_HEX] formats.
        token_pattern = re.compile(
            r"\[[A-Z]\d{5}\]"
            r"|\[[A-Z][A-Z0-9_]{4,40}\]"
        )
        found_tokens = set(token_pattern.findall(full_text))

        if not found_tokens:
            # No tokens found — return as-is
            return data, 0, []

        # Resolve each token via the vault
        token_map: dict[str, str] = {}  # token_string → original_text
        for token_str in found_tokens:
            mapping = vault.resolve_token(token_str)
            if mapping:
                token_map[token_str] = mapping.original_text
            else:
                all_unresolved.append(token_str)

        if not token_map:
            return data, 0, list(set(all_unresolved))

        # Second pass: erase tokens then insert original text per page
        for page in pdf_doc:
            # Collect deferred text insertions
            deferred: list[tuple[str, dict]] = []  # (original_text, style)

            for token_str, original_text in token_map.items():
                rects = page.search_for(token_str)
                if not rects:
                    continue

                for rect in rects:
                    # Extract the visual style of the token text
                    style = _extract_detok_style(page, rect)

                    # Widen erase rect to cover the original text which
                    # is usually longer than the compact token.
                    needed_width = fitz.get_text_length(
                        original_text,
                        fontname=style["fontname"],
                        fontsize=style["fontsize"],
                    )
                    erase_rect = fitz.Rect(rect)
                    if needed_width > erase_rect.width:
                        erase_rect.x1 = erase_rect.x0 + needed_width + 4

                    page.add_redact_annot(erase_rect, fill=(1, 1, 1))
                    deferred.append((original_text, style))
                    total_replaced += 1

            page.apply_redactions()

            # Insert original text at each token's baseline position
            for text, style in deferred:
                origin = style["origin"]
                page.insert_text(
                    fitz.Point(origin[0], origin[1]),
                    text,
                    fontname=style["fontname"],
                    fontsize=style["fontsize"],
                    color=style["text_color"],
                )

        output_bytes = pdf_doc.tobytes(deflate=True, clean=True)
        return output_bytes, total_replaced, list(set(all_unresolved))

    finally:
        pdf_doc.close()


def _extract_detok_style(page: fitz.Page, rect: fitz.Rect) -> dict:
    """Extract font properties from the dominant span overlapping *rect*.

    Returns dict with fontname (Base-14), fontsize, text_color, origin.
    Falls back to sensible defaults when no span overlaps.
    """
    from core.anonymizer.engine import _map_to_base14, _srgb_int_to_rgb

    blocks = page.get_text("dict", clip=rect.irect, flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

    best_span = None
    best_overlap = 0.0

    for block in blocks:
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                sb = fitz.Rect(span["bbox"])
                overlap = abs(sb & rect)  # intersection area
                if overlap > best_overlap:
                    best_overlap = overlap
                    best_span = span

    if best_span is None:
        return {
            "fontname": "helv",
            "fontsize": 11.0,
            "text_color": (0, 0, 0),
            "origin": (rect.x0, rect.y1 - 2),
        }

    span_origin = best_span.get("origin", (rect.x0, rect.y1 - 2))
    baseline_y = span_origin[1]

    return {
        "fontname": _map_to_base14(best_span.get("font", ""), best_span.get("flags", 0)),
        "fontsize": best_span.get("size", 11.0),
        "text_color": _srgb_int_to_rgb(best_span.get("color", 0)),
        "origin": (rect.x0, baseline_y),
    }


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

_HANDLERS: dict[str, Callable] = {
    ".txt":  _detokenize_txt,
    ".csv":  _detokenize_txt,
    ".docx": _detokenize_docx,
    ".xlsx": _detokenize_xlsx,
    ".pdf":  _detokenize_pdf,
}

SUPPORTED_EXTENSIONS = set(_HANDLERS.keys())


def detokenize_file(
    file_data: bytes,
    filename: str,
    vault,
) -> tuple[bytes, str, int, list[str]]:
    """
    De-tokenize tokens inside a file.

    Args:
        file_data: Raw file bytes.
        filename: Original filename (used to detect format).
        vault: Unlocked TokenVault instance.

    Returns:
        (output_bytes, output_filename, tokens_replaced, unresolved_tokens)
    """
    ext = _suffix_from_filename(filename)

    if ext == ".doc":
        raise ValueError(
            "Legacy .doc format is not supported. "
            "Please convert to .docx first (File → Save As in Word)."
        )
    if ext == ".xls":
        raise ValueError(
            "Legacy .xls format is not supported. "
            "Please convert to .xlsx first (File → Save As in Excel)."
        )

    handler = _HANDLERS.get(ext)
    if handler is None:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    output_data, count, unresolved = handler(file_data, vault)

    stem = Path(filename).stem
    out_name = f"{stem}_detokenized{ext}"

    logger.info(
        f"File de-tokenization: {filename} → {out_name}, "
        f"{count} token(s) replaced, {len(unresolved)} unresolved"
    )
    return output_data, out_name, count, unresolved
