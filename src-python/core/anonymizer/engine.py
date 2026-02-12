"""Anonymization engine — applies redaction and tokenization to documents."""

from __future__ import annotations

import asyncio
import io
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageFont

from core.config import config
from core.vault.store import vault
from models.schemas import (
    AnonymizeResponse,
    DocumentInfo,
    PIIRegion,
    PIIType,
    RegionAction,
    TokenMapping,
)

logger = logging.getLogger(__name__)

# File type constants
PDF_TYPES = {"application/pdf"}
IMAGE_TYPES = {"image/jpeg", "image/png", "image/tiff", "image/bmp", "image/webp"}
DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


# ---------------------------------------------------------------------------
# Shared finalization helpers (eliminate duplication across handlers)
# ---------------------------------------------------------------------------

def _save_encrypted_manifest(
    output_dir: Path,
    doc_id: str,
    original_filename: str,
    token_manifest: list[dict],
) -> None:
    """Save the token manifest encrypted with the vault key.

    The manifest maps tokens → original text.  Storing it in plaintext
    would undermine the vault encryption, so we encrypt it with the
    same Fernet key.
    """
    manifest_data = json.dumps({
        "doc_id": doc_id,
        "original_filename": original_filename,
        "tokens": token_manifest,
    }).encode("utf-8")

    # Encrypt using the vault's Fernet instance
    if vault._fernet is not None:
        encrypted = vault._fernet.encrypt(manifest_data)
        manifest_path = output_dir / "token_manifest.enc"
        manifest_path.write_bytes(encrypted)
    else:
        # Fallback: save only token strings (no originals) if vault key is unavailable
        safe_manifest = {
            "doc_id": doc_id,
            "original_filename": original_filename,
            "tokens": [{"token_string": t["token_string"]} for t in token_manifest],
        }
        manifest_path = output_dir / "token_manifest.json"
        manifest_path.write_text(json.dumps(safe_manifest, indent=2), encoding="utf-8")

    logger.info(f"Saved token manifest with {len(token_manifest)} entries (encrypted)")


def _finalize_anonymization(
    doc: DocumentInfo,
    output_path: Path,
    tokens_created: int,
    regions_removed: int,
    token_manifest: list[dict],
) -> AnonymizeResponse:
    """Common finalization: save manifest, register document, return response."""
    output_dir = output_path.parent

    # Save encrypted manifest
    _save_encrypted_manifest(output_dir, doc.doc_id, doc.original_filename, token_manifest)

    # Register in vault
    vault.register_document(
        doc_id=doc.doc_id,
        original_filename=doc.original_filename,
        page_count=doc.page_count,
        anonymized_filename=output_path.name,
    )

    logger.info(
        f"Anonymized '{doc.original_filename}': "
        f"{regions_removed} removed, {tokens_created} tokenized"
    )

    return AnonymizeResponse(
        doc_id=doc.doc_id,
        output_path=str(output_path),
        output_text_path="",
        tokens_created=tokens_created,
        regions_removed=regions_removed,
    )


async def anonymize_document(doc: DocumentInfo) -> AnonymizeResponse:
    """
    Apply all accepted region actions to a document and produce output files.

    Dispatches to the appropriate handler based on file type:
    - PDF: Text-based redaction with PyMuPDF
    - DOCX: Text replacement in paragraphs/tables
    - XLSX: Text replacement in cells
    - Images: Bitmap manipulation
    """
    if not vault.is_unlocked:
        raise RuntimeError("Vault must be unlocked before anonymization")

    original_path = Path(doc.file_path)
    if not original_path.exists():
        raise RuntimeError(f"Original file not found: {original_path}")

    # Dispatch to appropriate handler based on MIME type
    if doc.mime_type in PDF_TYPES:
        return await asyncio.to_thread(_anonymize_pdf_sync, doc, original_path)
    elif doc.mime_type == DOCX_TYPE:
        return await asyncio.to_thread(_anonymize_docx_sync, doc, original_path)
    elif doc.mime_type == XLSX_TYPE:
        return await asyncio.to_thread(_anonymize_xlsx_sync, doc, original_path)
    elif doc.mime_type in IMAGE_TYPES:
        return await asyncio.to_thread(_anonymize_image_sync, doc, original_path)
    else:
        raise ValueError(f"Unsupported file type for anonymization: {doc.mime_type}")


def _anonymize_pdf_sync(doc: DocumentInfo, original_path: Path) -> AnonymizeResponse:
    """Anonymize a PDF file using text-based redactions."""
    tokens_created = 0
    regions_removed = 0
    token_manifest: list[dict] = []

    # Open PDF with PyMuPDF
    pdf_doc = fitz.open(str(original_path))

    try:
        # Widen the redaction rect so the (longer) token text fits at the
        # original font size.  PyMuPDF auto-shrinks replacement text when
        # it doesn't fit; by pre-extending the rect we avoid that.

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            page_regions = [
                r for r in doc.regions if r.page_number == page_num + 1
            ]

            for region in page_regions:
                bbox = region.bbox
                rect = fitz.Rect(bbox.x0, bbox.y0, bbox.x1, bbox.y1)

                if region.action == RegionAction.REMOVE:
                    replacement_text = "---"
                elif region.action == RegionAction.TOKENIZE:
                    token_string = vault.generate_token_string(region.pii_type)

                    mapping = TokenMapping(
                        token_string=token_string,
                        original_text=region.text,
                        pii_type=region.pii_type,
                        source_document=doc.original_filename,
                        context_snippet=_get_context_snippet(
                            doc.pages[page_num].full_text,
                            region.char_start,
                            region.char_end,
                        ),
                    )
                    vault.store_token(mapping)
                    replacement_text = token_string

                    token_manifest.append({
                        "token_string": token_string,
                        "original_text": region.text,
                        "page_number": page_num + 1,
                    })
                else:
                    continue

                # Extract the original visual style before we redact
                style = _extract_span_style(page, rect)

                # Measure replacement text width at the original font size
                # and extend the rect if the token string is wider.
                needed_width = fitz.get_text_length(
                    replacement_text,
                    fontname=style["fontname"],
                    fontsize=style["fontsize"],
                )
                current_width = rect.width
                if needed_width > current_width:
                    rect.x1 = rect.x0 + needed_width + 2  # +2pt padding

                page.add_redact_annot(
                    rect,
                    text=replacement_text,
                    fontname=style["fontname"],
                    fontsize=style["fontsize"],
                    fill=(1, 1, 1),
                    text_color=style["text_color"],
                )

                if region.action == RegionAction.REMOVE:
                    regions_removed += 1
                else:
                    tokens_created += 1

            # Apply all redactions on this page
            page.apply_redactions()

        # ── Metadata scrubbing ─────────────────────────────────────────
        # 1. Clear standard document metadata (Author, Title, Subject, …)
        pdf_doc.set_metadata({})

        # 2. Strip XMP (XML-based) metadata
        pdf_doc.del_xml_metadata()

        # 3. Remove table-of-contents / bookmarks (may contain PII names)
        pdf_doc.set_toc([])

        # 4. Remove all remaining annotations (comments, highlights,
        #    sticky notes, etc.) — redaction annotations were already
        #    consumed by apply_redactions(), but other kinds may survive.
        for pg_idx in range(len(pdf_doc)):
            pg = pdf_doc[pg_idx]
            annot_list = list(pg.annots()) if pg.annots() else []
            for annot in annot_list:
                pg.delete_annot(annot)

        # 5. Remove embedded file attachments (portfolio / attached docs)
        try:
            if pdf_doc.embfile_count() > 0:
                names = [pdf_doc.embfile_info(i)["name"]
                         for i in range(pdf_doc.embfile_count())]
                for name in names:
                    pdf_doc.embfile_del(name)
        except Exception:
            logger.debug("Could not enumerate/remove embedded files")

        logger.info("Scrubbed PDF metadata, annotations, TOC, and attachments")

        # Save anonymized PDF
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        stem = Path(doc.original_filename).stem
        output_dir = config.temp_dir / doc.doc_id / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        pdf_path = output_dir / f"{stem}_anonymized_{timestamp}.pdf"
        pdf_doc.save(str(pdf_path), deflate=True, clean=True)
        logger.info(f"Saved anonymized PDF: {pdf_path}")

        return _finalize_anonymization(
            doc, pdf_path, tokens_created, regions_removed, token_manifest,
        )

    finally:
        pdf_doc.close()


def _replace_in_paragraphs(paragraphs, replacements: dict[str, str]) -> None:
    """Replace text in DOCX paragraphs while preserving per-run formatting.

    For each paragraph:
    1. Join all run texts into a single string and build a character→run map.
    2. Apply replacements on the joined text, tracking where each
       replacement lands.
    3. Redistribute the result back into the original runs so only the
       runs that overlap a replacement are touched — all other runs keep
       their formatting and text verbatim.

    This handles PII that is split across multiple XML runs
    (e.g., "Joh" | "n S" | "mith") while preserving styles everywhere else.
    """
    for paragraph in paragraphs:
        runs = paragraph.runs
        if not runs:
            continue
        # Build original text and a list of (start, end) offsets per run
        run_texts = [r.text or "" for r in runs]
        full = "".join(run_texts)
        if not full:
            continue

        # Apply all replacements on the joined text
        replaced = full
        for original, replacement in replacements.items():
            replaced = replaced.replace(original, replacement)
        if replaced == full:
            continue

        # Build character→run_index mapping for the ORIGINAL text.
        # We process replacements one at a time, adjusting offsets.
        # The strategy: find each replacement span in `full`, note which
        # runs it overlaps, push the replacement into the first overlapping
        # run and trim text from the others.
        #
        # We work on a mutable list of run-text strings so we can do
        # multiple passes (one per replacement) without losing track.
        new_run_texts = list(run_texts)

        for original, replacement in replacements.items():
            # Repeatedly replace all occurrences in the joined run texts
            while True:
                # Re-join current state
                joined = "".join(new_run_texts)
                idx = joined.find(original)
                if idx == -1:
                    break
                end = idx + len(original)

                # Determine which runs are affected
                cursor = 0
                for ri, rt in enumerate(new_run_texts):
                    run_start = cursor
                    run_end = cursor + len(rt)

                    if run_end <= idx:
                        # Entirely before the match
                        cursor = run_end
                        continue
                    if run_start >= end:
                        # Entirely after the match — done
                        break

                    # This run overlaps the match
                    local_start = max(idx - run_start, 0)
                    local_end = min(end - run_start, len(rt))

                    if run_start <= idx < run_end:
                        # First overlapping run — insert replacement here
                        new_run_texts[ri] = rt[:local_start] + replacement + rt[local_end:]
                    else:
                        # Subsequent overlapping runs — just remove the matched portion
                        new_run_texts[ri] = rt[:local_start] + rt[local_end:]

                    cursor = run_end

        # Write back to the actual run objects
        for ri, run in enumerate(runs):
            run.text = new_run_texts[ri]


def _replace_in_docx_xml_parts(docx_doc, replacements: dict[str, str]) -> None:
    """Replace PII in DOCX XML parts that python-docx doesn't expose natively.

    Covers footnotes, endnotes, and any custom XML parts that might
    contain PII text.  Works by directly manipulating the underlying
    XML of each relevant part in the OPC package.
    """
    import re
    from lxml import etree

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Part URIs that may contain user text with PII
    text_part_uris = [
        "/word/footnotes.xml",
        "/word/endnotes.xml",
        "/word/comments.xml",     # also nuke comments
    ]

    for part_uri in text_part_uris:
        try:
            part = docx_doc.part.package.part_related_by(part_uri)
        except Exception:
            # Part may not exist in this document
            try:
                # Fallback: iterate all parts and match by partname
                part = None
                for p in docx_doc.part.package.iter_parts():
                    if str(p.partname) == part_uri:
                        part = p
                        break
                if part is None:
                    continue
            except Exception:
                continue

        try:
            xml_bytes = part.blob
            root = etree.fromstring(xml_bytes)

            # Replace text in all <w:t> elements
            changed = False
            for t_elem in root.iter(f"{{{W_NS}}}t"):
                if t_elem.text:
                    new_text = t_elem.text
                    for original, replacement in replacements.items():
                        new_text = new_text.replace(original, replacement)
                    if new_text != t_elem.text:
                        t_elem.text = new_text
                        changed = True

            if changed:
                part._blob = etree.tostring(root, xml_declaration=True,
                                             encoding="UTF-8", standalone=True)
                logger.debug(f"Replaced PII in DOCX part: {part_uri}")
        except Exception:
            logger.debug(f"Could not process DOCX part: {part_uri}")


def _anonymize_docx_sync(doc: DocumentInfo, original_path: Path) -> AnonymizeResponse:
    """Anonymize a DOCX file using text replacements."""
    from docx import Document
    
    tokens_created = 0
    regions_removed = 0
    token_manifest: list[dict] = []

    # Open DOCX
    docx = Document(str(original_path))

    try:
        # Sort regions by char_end descending to preserve offsets during replacement
        sorted_regions = sorted(
            [r for r in doc.regions if r.action in (RegionAction.REMOVE, RegionAction.TOKENIZE)],
            key=lambda r: r.char_end,
            reverse=True,
        )

        # Build full text to find replacements
        full_text = "\n".join(p.full_text for p in doc.pages) if doc.pages else ""

        # Track replacements
        replacements: dict[str, str] = {}

        for region in sorted_regions:
            original_text = region.text

            if region.action == RegionAction.REMOVE:
                replacement = "---"
                replacements[original_text] = replacement
                regions_removed += 1

            elif region.action == RegionAction.TOKENIZE:
                token_string = vault.generate_token_string(region.pii_type)

                mapping = TokenMapping(
                    token_string=token_string,
                    original_text=original_text,
                    pii_type=region.pii_type,
                    source_document=doc.original_filename,
                    context_snippet=_get_context_snippet(full_text, region.char_start, region.char_end),
                )
                vault.store_token(mapping)

                replacements[original_text] = token_string
                tokens_created += 1

                token_manifest.append({
                    "token_string": token_string,
                    "original_text": original_text,
                })

        # Apply replacements to all text in document
        _replace_in_paragraphs(docx.paragraphs, replacements)

        # Apply to tables
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    _replace_in_paragraphs(cell.paragraphs, replacements)

        # Apply to headers and footers (PII often appears here)
        for section in docx.sections:
            for hf in (section.header, section.footer,
                        section.first_page_header, section.first_page_footer,
                        section.even_page_header, section.even_page_footer):
                if hf is None:
                    continue
                _replace_in_paragraphs(hf.paragraphs, replacements)
                for tbl in hf.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            _replace_in_paragraphs(cell.paragraphs, replacements)

        # Apply to footnotes and endnotes via underlying XML
        _replace_in_docx_xml_parts(docx, replacements)

        # ── Metadata scrubbing ─────────────────────────────────────────
        props = docx.core_properties
        props.author = ""
        props.last_modified_by = ""
        props.title = ""
        props.subject = ""
        props.keywords = ""
        props.comments = ""
        props.category = ""
        logger.info("Scrubbed DOCX metadata, headers, footers, and footnotes")

        # Save anonymized DOCX
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        stem = Path(doc.original_filename).stem
        output_dir = config.temp_dir / doc.doc_id / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        docx_path = output_dir / f"{stem}_anonymized_{timestamp}.docx"
        docx.save(str(docx_path))
        logger.info(f"Saved anonymized DOCX: {docx_path}")

        return _finalize_anonymization(
            doc, docx_path, tokens_created, regions_removed, token_manifest,
        )

    finally:
        # python-docx doesn't have a close() but we ensure no dangling refs
        del docx


def _anonymize_xlsx_sync(doc: DocumentInfo, original_path: Path) -> AnonymizeResponse:
    """Anonymize an XLSX file using text replacements in cells."""
    from openpyxl import load_workbook

    tokens_created = 0
    regions_removed = 0
    token_manifest: list[dict] = []

    # Open XLSX
    wb = load_workbook(str(original_path))

    try:
        # Build full text to match regions
        full_text = "\n".join(p.full_text for p in doc.pages) if doc.pages else ""

        # Sort regions by char_end descending
        sorted_regions = sorted(
            [r for r in doc.regions if r.action in (RegionAction.REMOVE, RegionAction.TOKENIZE)],
            key=lambda r: r.char_end,
            reverse=True,
        )

        # Track replacements
        replacements: dict[str, str] = {}

        for region in sorted_regions:
            original_text = region.text

            if region.action == RegionAction.REMOVE:
                replacement = "---"
                replacements[original_text] = replacement
                regions_removed += 1

            elif region.action == RegionAction.TOKENIZE:
                token_string = vault.generate_token_string(region.pii_type)

                mapping = TokenMapping(
                    token_string=token_string,
                    original_text=original_text,
                    pii_type=region.pii_type,
                    source_document=doc.original_filename,
                    context_snippet=_get_context_snippet(full_text, region.char_start, region.char_end),
                )
                vault.store_token(mapping)

                replacements[original_text] = token_string
                tokens_created += 1

                token_manifest.append({
                    "token_string": token_string,
                    "original_text": original_text,
                })

        # Apply replacements to all cells
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str) and cell.value:
                        for original, replacement in replacements.items():
                            cell.value = cell.value.replace(original, replacement)

        # Also check cell comments / notes for PII
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.comment and cell.comment.text:
                        for original, replacement in replacements.items():
                            cell.comment.text = cell.comment.text.replace(
                                original, replacement
                            )

        # ── Metadata scrubbing ─────────────────────────────────────────
        wb.properties.creator = ""
        wb.properties.lastModifiedBy = ""
        wb.properties.title = ""
        wb.properties.subject = ""
        wb.properties.description = ""
        wb.properties.keywords = ""
        wb.properties.category = ""
        logger.info("Scrubbed XLSX metadata and cell comments")

        # Save anonymized XLSX
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        stem = Path(doc.original_filename).stem
        output_dir = config.temp_dir / doc.doc_id / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        xlsx_path = output_dir / f"{stem}_anonymized_{timestamp}.xlsx"
        wb.save(str(xlsx_path))
        logger.info(f"Saved anonymized XLSX: {xlsx_path}")

        return _finalize_anonymization(
            doc, xlsx_path, tokens_created, regions_removed, token_manifest,
        )

    finally:
        wb.close()


def _anonymize_image_sync(doc: DocumentInfo, original_path: Path) -> AnonymizeResponse:
    """Anonymize an image file using bitmap manipulation."""
    tokens_created = 0
    regions_removed = 0
    token_manifest: list[dict] = []

    # Load the original image
    img = Image.open(str(original_path)).convert("RGB")

    try:
        draw = ImageDraw.Draw(img)

        # Get page dimensions
        page_data = doc.pages[0] if doc.pages else None
        if not page_data:
            raise ValueError("No page data available for image")

        full_text = page_data.full_text

        # Process all annotated regions
        for region in doc.regions:
            if region.action == RegionAction.REMOVE:
                # Draw white rectangle (clean removal)
                bbox = region.bbox
                draw.rectangle(
                    [bbox.x0, bbox.y0, bbox.x1, bbox.y1],
                    fill=(255, 255, 255),
                )
                regions_removed += 1

            elif region.action == RegionAction.TOKENIZE:
                token_string = vault.generate_token_string(region.pii_type)

                mapping = TokenMapping(
                    token_string=token_string,
                    original_text=region.text,
                    pii_type=region.pii_type,
                    source_document=doc.original_filename,
                    context_snippet=_get_context_snippet(full_text, region.char_start, region.char_end),
                )
                vault.store_token(mapping)

                # Draw white rectangle then add token text
                bbox = region.bbox
                draw.rectangle(
                    [bbox.x0, bbox.y0, bbox.x1, bbox.y1],
                    fill=(255, 255, 255),
                )
                # Draw token text sized to match the original region height
                region_h = bbox.y1 - bbox.y0
                target_size = max(8, int(region_h * 0.75))
                font = None
                try:
                    font = ImageFont.truetype("arial.ttf", target_size)
                except (OSError, IOError):
                    try:
                        font = ImageFont.truetype("DejaVuSans.ttf", target_size)
                    except (OSError, IOError):
                        font = ImageFont.load_default()
                draw.text(
                    (bbox.x0 + 2, bbox.y0 + (region_h - target_size) / 2),
                    token_string,
                    fill=(0, 0, 0),
                    font=font,
                )
                tokens_created += 1

                token_manifest.append({
                    "token_string": token_string,
                    "original_text": region.text,
                })

        # Save anonymized image with original extension
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        stem = Path(doc.original_filename).stem
        ext = Path(doc.original_filename).suffix.lower()
        output_dir = config.temp_dir / doc.doc_id / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        img_path = output_dir / f"{stem}_anonymized_{timestamp}{ext}"

        # ── Metadata scrubbing ─────────────────────────────────────────
        # Strip ALL metadata (EXIF, IPTC, XMP, ICC profiles, PNG text chunks)
        # that may contain PII such as GPS coords, author, camera serial, etc.
        # Create a clean pixel-only copy of the image.
        clean_img = Image.new(img.mode, img.size)
        clean_img.putdata(list(img.getdata()))

        # Save with appropriate format
        try:
            if ext in [".jpg", ".jpeg"]:
                clean_img.save(str(img_path), "JPEG", quality=95)
            elif ext == ".png":
                clean_img.save(str(img_path), "PNG")
            else:
                clean_img.save(str(img_path))
        finally:
            clean_img.close()

        logger.info(f"Saved anonymized image (EXIF stripped): {img_path}")

        return _finalize_anonymization(
            doc, img_path, tokens_created, regions_removed, token_manifest,
        )

    finally:
        img.close()


def _map_to_base14(font_name: str, flags: int) -> str:
    """Map an arbitrary PDF font name + flags to the closest Base-14 font.

    ``flags`` comes from a PyMuPDF span dict (bit 1 = italic, bit 4 = bold,
    bit 2 = serif, bit 3 = monospaced).
    """
    is_bold = bool(flags & 16)
    is_italic = bool(flags & 2)
    is_serif = bool(flags & 4)
    is_mono = bool(flags & 8)

    name_lower = font_name.lower()

    # Monospace family
    if is_mono or any(k in name_lower for k in ("courier", "mono", "consol", "firacode", "source code")):
        if is_bold and is_italic:
            return "cobi"
        if is_bold:
            return "cobo"
        if is_italic:
            return "coit"
        return "cour"

    # Serif family
    if is_serif or any(k in name_lower for k in ("times", "serif", "roman", "garamond", "georgia", "cambria", "palat")):
        if is_bold and is_italic:
            return "tibi"
        if is_bold:
            return "tibo"
        if is_italic:
            return "tiit"
        return "tiro"

    # Default: Helvetica (sans-serif)
    if is_bold and is_italic:
        return "hebi"
    if is_bold:
        return "hebo"
    if is_italic:
        return "heit"
    return "helv"


def _srgb_int_to_rgb(color_int: int) -> tuple[float, float, float]:
    """Convert a span ``color`` integer (0xRRGGBB) to the (r, g, b) float
    tuple that PyMuPDF drawing/redaction methods expect (each 0.0–1.0)."""
    r = ((color_int >> 16) & 0xFF) / 255.0
    g = ((color_int >> 8) & 0xFF) / 255.0
    b = (color_int & 0xFF) / 255.0
    return (r, g, b)


def _extract_span_style(
    page: fitz.Page,
    rect: fitz.Rect,
) -> dict:
    """Return the dominant text-span style properties inside *rect*.

    Searches all text spans that overlap the redaction rectangle and picks
    the one with the largest overlap area so the replacement inherits the
    correct font size, weight, and color.

    Also returns the text *origin* (baseline insertion point) so that
    replacement text can be inserted at exactly the right position.

    Falls back to sensible defaults if the rect contains no extractable text
    (e.g. image-only area).
    """
    blocks = page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT).get("blocks", [])

    best_span: dict | None = None
    best_overlap: float = 0.0

    for block in blocks:
        if block.get("type", 0) != 0:  # text blocks only
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                span_rect = fitz.Rect(span["bbox"])
                intersection = span_rect & rect
                if intersection.is_empty:
                    continue
                area = intersection.width * intersection.height
                if area > best_overlap:
                    best_overlap = area
                    best_span = span

    if best_span is None:
        return {
            "fontname": "helv",
            "fontsize": 11.0,
            "text_color": (0, 0, 0),
            "origin": (rect.x0, rect.y1 - 2),  # rough baseline fallback
        }

    return {
        "fontname": _map_to_base14(best_span.get("font", ""), best_span.get("flags", 0)),
        "fontsize": best_span.get("size", 11.0),
        "text_color": _srgb_int_to_rgb(best_span.get("color", 0)),
        "origin": tuple(best_span.get("origin", (rect.x0, rect.y1 - 2))),
    }


def _get_context_snippet(text: str, start: int, end: int, context_chars: int = 50) -> str:
    """Extract a text snippet around the PII for disambiguation."""
    ctx_start = max(0, start - context_chars)
    ctx_end = min(len(text), end + context_chars)
    snippet = text[ctx_start:ctx_end]
    if ctx_start > 0:
        snippet = "..." + snippet
    if ctx_end < len(text):
        snippet = snippet + "..."
    return snippet
